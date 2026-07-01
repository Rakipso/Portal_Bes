import os
import subprocess
import sys

def install(package):
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', package])

try:
    import docx
except ImportError:
    install('python-docx')
    import docx

try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
except ImportError:
    install('openpyxl')
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

# --- 1. MODIFICAR EL DOCX ORIGINAL DIRECTAMENTE ---
docx_path = r"c:\Users\teamf\Downloads\proyecto Portal-bes\portal-migracion\Plan_de_Pruebas_Portal_BES_v1.docx"

try:
    # Abrimos el documento existente en lugar de crear uno nuevo
    doc = docx.Document(docx_path)
    
    # Agregamos el contenido faltante al final del documento conservando el formato
    doc.add_heading('4.1 a 4.5 Resumen de Cobertura Normativa', level=2)
    doc.add_paragraph('De acuerdo con el modelo de calidad definido por la norma ISO 9126 (evaluando Funcionalidad, Fiabilidad y Usabilidad) y complementado con la norma moderna ISO/IEC 25010 (Seguridad, Accesibilidad y Rendimiento), se han diseñado 18 casos de prueba estructurados, excediendo el mínimo exigido para el equipo.')
    doc.add_paragraph('HU-01 a HU-02 (Funcionalidad y Fiabilidad ISO 9126): Validados a través de los casos TC_001 al TC_006, garantizando que el software entrega las funciones requeridas y mantiene su nivel de desempeño en casos de caídas de APIs (tolerancia a fallos).')
    doc.add_paragraph('HU-03 (Usabilidad ISO 9126): El entendimiento claro de los motivos de rechazo se cubrió en los casos TC_007 al TC_009.')
    doc.add_paragraph('HU-04 (Seguridad ISO 25000): La confidencialidad y limpieza de estado de sesión se evaluó en los casos TC_010 a TC_012.')
    doc.add_paragraph('HU-05 (Accesibilidad ISO 25000): Garantizada mediante pruebas con teclados y lectores NVDA (TC_013 a TC_015).')
    
    doc.add_heading('5. Trazabilidad de Defectos (Bugs)', level=1)
    doc.add_paragraph('Durante la simulación de ejecución en el tablero Kanban, los casos TC_007, TC_010, TC_015 y TC_017 resultaron con estado "Failed". Esto derivó en el levantamiento de los defectos BUG_001 a BUG_004 en nuestro reporte de hallazgos. Posteriormente, los errores críticos (ej. BUG_004 que impedía entrar al panel admin) fueron resueltos e integrados al Control de Cambios en la versión v1.2.')
    
    doc.add_heading('6. Control de Cambios y Versionamiento', level=1)
    doc.add_paragraph('El control de versiones se manejó de manera exhaustiva documentando cómo el prototipo evolucionó desde su estado conceptual (v1.0) hasta el prototipo robusto (v1.4) mitigando deficiencias de servidor (Cold Starts en Vercel) para asegurar la Fiabilidad y Usabilidad (ISO 9126). El desglose completo se adjunta en el Excel de Entregables.')
    
    # Sobrescribimos el mismo archivo
    doc.save(docx_path)
    print("DOCX original modificado correctamente.")
except Exception as e:
    print(f"Error modificando DOCX: {e}")

# --- 2. DAR FORMATO PRESENTABLE AL EXCEL ---
excel_path = r"c:\Users\teamf\Downloads\proyecto Portal-bes\portal-migracion\Entregable_Calidad_Portal_BES.xlsx"

try:
    wb = openpyxl.load_workbook(excel_path)
    
    # Estilos
    header_fill = PatternFill(start_color="0F4C81", end_color="0F4C81", fill_type="solid") # Azul Portal BES
    header_font = Font(color="FFFFFF", bold=True)
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
    alignment = Alignment(wrap_text=True, vertical='center')
    
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        
        # Estilizar la primera fila (Cabeceras)
        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = thin_border
            
        # Aplicar bordes, alineación y auto-ajustar columnas
        for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=ws.max_column):
            for cell in row:
                cell.border = thin_border
                cell.alignment = alignment
        
        # Ajustar anchos de columna (aproximación)
        for col in ws.columns:
            max_length = 0
            column = col[0].column_letter
            for cell in col:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = (max_length + 2)
            if adjusted_width > 40:
                adjusted_width = 40 # Límite máximo para que no sea excesivo, el wrap_text hará el resto
            ws.column_dimensions[column].width = adjusted_width

    wb.save(excel_path)
    print("Excel formateado correctamente.")
except Exception as e:
    print(f"Error formateando Excel: {e}")
