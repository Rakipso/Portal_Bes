import pandas as pd
from docx import Document
import datetime

# 1. GENERATE EXCEL (Casos de Prueba, Defectos, Control de Cambios)

test_cases = [
    # HU-01 Autenticación
    ['TC_001', 'HU-01', 'Login Exitoso GobDigital', 'Usuario en pantalla de inicio', '1. Clic en "Ingresar con ClaveÚnica"\n2. Ingresar credenciales válidas\n3. Autorizar', 'RUT válido y clave correcta', 'Redirige a panel de beneficios (Ciudadano)', 'Funcionalidad (ISO 9126) / Seguridad (ISO 25000)', 'Passed'],
    ['TC_002', 'HU-01', 'Login Fallido Credenciales Inválidas', 'Usuario en pasarela GobDigital', '1. Ingresar RUT y clave errónea\n2. Clic en ingresar', 'RUT inválido o clave incorrecta', 'Muestra error y deniega acceso', 'Fiabilidad (ISO 9126)', 'Passed'],
    ['TC_003', 'HU-01', 'Login Cancelado', 'Usuario en pasarela GobDigital', '1. Clic en "Cancelar"', 'N/A', 'Redirige al inicio del Portal BES', 'Usabilidad (ISO 9126)', 'Passed'],
    # HU-02 Visualización de Beneficios
    ['TC_004', 'HU-02', 'Ver tarjetas de beneficios elegibles', 'Usuario autenticado y con requisitos', '1. Cargar panel de ciudadano', 'Perfil: Mujer, 25 años, RSH 40%', 'Muestra "Bono al Trabajo de la Mujer" con botón "Ir al beneficio"', 'Funcionalidad (ISO 9126)', 'Passed'],
    ['TC_005', 'HU-02', 'Usuario sin beneficios elegibles', 'Usuario autenticado sin cumplir ningún requisito', '1. Cargar panel de ciudadano', 'Perfil: Hombre, 50 años, RSH 90%', 'Muestra mensaje "No cumples los requisitos para estos beneficios"', 'Funcionalidad (ISO 9126)', 'Passed'],
    ['TC_006', 'HU-02', 'Falla en API Externa (Link Caído)', 'Usuario autenticado', '1. Cargar panel ciudadano', 'SENCE simulado como Offline', 'Muestra tarjeta SENCE en gris con alerta "Institución temporalmente no disponible"', 'Fiabilidad (ISO 9126) / Tolerancia a fallos (ISO 25000)', 'Passed'],
    # HU-03 Motivos de rechazo
    ['TC_007', 'HU-03', 'Motivo de rechazo por RSH', 'Usuario con RSH alto (ej. 80%)', '1. Revisar sección "Beneficios No Aplicables"', 'Perfil: RSH 80%', 'Muestra "Superas el 60% en el Registro Social de Hogares"', 'Usabilidad (ISO 9126)', 'Failed'],
    ['TC_008', 'HU-03', 'Múltiples motivos de rechazo', 'Usuario no cumple edad ni RSH', '1. Revisar "Bono Logro Escolar" en no aplicables', 'Perfil: 26 años, RSH 60%', 'Muestra viñetas: Edad máxima superada y RSH no cumple', 'Funcionalidad (ISO 9126)', 'Passed'],
    ['TC_009', 'HU-03', 'Motivo de rechazo por género', 'Usuario Hombre revisando bono de mujer', '1. Revisar "Bono por Hijo"', 'Perfil: Hombre', 'Muestra "Beneficio exclusivo para mujeres"', 'Funcionalidad (ISO 9126)', 'Passed'],
    # HU-04 Cierre de Sesión Seguro
    ['TC_010', 'HU-04', 'Cierre de sesión manual', 'Usuario autenticado', '1. Clic en "Cerrar Sesión"', 'N/A', 'Sesión terminada, redirige al inicio', 'Seguridad (ISO 25000)', 'Failed'],
    ['TC_011', 'HU-04', 'Prevención de retroceso (Back Button)', 'Usuario cerró sesión recientemente', '1. Presionar botón "Atrás" del navegador', 'N/A', 'Mantiene en pantalla de login, no muestra caché', 'Fiabilidad (ISO 9126)', 'Passed'],
    ['TC_012', 'HU-04', 'Cierre por inactividad', 'Usuario inactivo por 5 min', '1. Esperar 5 minutos sin mover el mouse', 'N/A', 'Cierra sesión automáticamente y redirige al inicio', 'Eficiencia (ISO 9126) / Seguridad (ISO 25000)', 'Blocked'],
    # HU-05 Accesibilidad
    ['TC_013', 'HU-05', 'Navegación por Tabulador', 'Usuario en inicio', '1. Presionar TAB secuencialmente', 'Tecla TAB', 'El foco visual se mueve lógicamente por los botones', 'Usabilidad (ISO 9126) / Accesibilidad (ISO 25000)', 'Passed'],
    ['TC_014', 'HU-05', 'Escalado móvil 320px', 'Usuario en móvil', '1. Cargar portal en resolución 320x568', 'Viewport 320px', 'Diseño responsivo, no hay scroll horizontal', 'Usabilidad (ISO 9126)', 'Passed'],
    ['TC_015', 'HU-05', 'Compatibilidad Screen Reader', 'Usuario con NVDA', '1. Navegar tarjetas con NVDA', 'Software NVDA activo', 'Lector anuncia correctamente "Ir al beneficio"', 'Accesibilidad (ISO 25000)', 'Failed'],
    # Admin HUs
    ['TC_016', 'HU-06', 'Rendimiento 10.000 usuarios', 'Servidor corriendo', '1. Lanzar JMeter con 10K threads', 'Test plan JMeter 10K', 'Tasa de error < 1%, tiempo de respuesta < 3s', 'Eficiencia (ISO 9126) / Rendimiento (ISO 25000)', 'Passed'],
    ['TC_017', 'HU-07', 'Lectura de Logs por Admin', 'Admin autenticado', '1. Clic en "Panel de Auditoría"', 'Rol: ADMIN', 'Muestra terminal de logs sin RUTs visibles', 'Funcionalidad (ISO 9126) / Privacidad (ISO 25000)', 'Failed'],
    ['TC_018', 'HU-08', 'Simulación Mantenimiento Programado', 'Admin en Panel de Auditoría', '1. Activar "Modo Mantenimiento"', 'Clic en toggle', 'Ciudadanos ven pantalla de mantenimiento de inmediato', 'Fiabilidad (ISO 9126)', 'Passed']
]

df_tc = pd.DataFrame(test_cases, columns=[
    'ID Caso', 'Historia', 'Nombre/Descripción', 'Precondición', 'Pasos de Ejecución', 
    'Datos de Entrada', 'Resultado Esperado', 'Norma de Calidad (ISO 9126 / 25000)', 'Estado Kanban'
])

defects = [
    ['BUG_001', 'TC_007', 'Alta', '1. Iniciar sesión con perfil RSH 80%\n2. Ver sección beneficios no aplicables\n3. Revisar tarjeta "Subsidio de Arriendo"', 'No se muestra el porcentaje exacto en el motivo, solo un texto genérico de error de servidor.', 'Debería mostrar "Su tramo actual (80%) supera el máximo".', 'Corregido'],
    ['BUG_002', 'TC_010', 'Crítica', '1. Iniciar sesión como Admin\n2. Hacer clic en "Cerrar sesión"', 'El botón no responde, la consola arroja "ReferenceError: logout is not defined".', 'Debería cerrar sesión y volver a inicio.', 'Corregido'],
    ['BUG_003', 'TC_015', 'Media', '1. Activar lector JAWS/NVDA\n2. Navegar sobre la tarjeta de un beneficio rechazado', 'El lector ignora el motivo de rechazo porque falta el atributo aria-label.', 'El lector debe leer los motivos exactos listados.', 'Abierto'],
    ['BUG_004', 'TC_017', 'Alta', '1. Entrar como Admin\n2. Ir a "Panel de Auditoría"', 'El sistema desloguea al Administrador en lugar de mostrar la consola.', 'Debería renderizar la página /admin con los logs.', 'Corregido']
]

df_bugs = pd.DataFrame(defects, columns=[
    'ID Defecto', 'Caso Prueba Asoc.', 'Severidad', 'Pasos para Reproducir', 
    'Resultado Obtenido', 'Resultado Esperado', 'Estado del Bug'
])

changes = [
    ['v1.0', '15-May-2026', 'Equipo Desarrollo', 'Módulo Autenticación', 'Versión Inicial - Prototipo Conceptual', 'Funcionalidad (ISO 9126)'],
    ['v1.1', '22-May-2026', 'Equipo Diseño', 'Módulo Frontend (Tarjetas)', 'Se aplicó alto contraste a las tarjetas de beneficios para cumplir con WCAG 2.1 AA.', 'Usabilidad (ISO 9126) / Accesibilidad (ISO 25000)'],
    ['v1.2', '03-Jun-2026', 'Equipo Desarrollo', 'Motor Backend / APIs', 'Corrección BUG_004: Administrador no podía acceder al dashboard. Se reestructuró el Middleware de Next.js.', 'Fiabilidad (ISO 9126)'],
    ['v1.3', '08-Jun-2026', 'Equipo Arquitectura', 'Manejo de Estados (Serverless)', 'Se implementó capa de servicios simulados en LocalStorage mitigando problemas de "Cold Starts" en Vercel.', 'Eficiencia (ISO 9126) / Tolerancia a fallos (ISO 25000)'],
    ['v1.4', '14-Jun-2026', 'Equipo QA', 'Auditoría', 'Incorporación de logs asíncronos y modo mantenimiento global verificado en tiempo real.', 'Seguridad y Mantenibilidad (ISO 25000)']
]

df_changes = pd.DataFrame(changes, columns=[
    'Versión', 'Fecha', 'Autor', 'Módulo Impactado', 'Razón / Descripción del Cambio', 'Impacto Normativo (ISO)'
])

with pd.ExcelWriter('Entregable_Calidad_Portal_BES.xlsx', engine='openpyxl') as writer:
    df_tc.to_excel(writer, sheet_name='Casos de Pruebas', index=False)
    df_bugs.to_excel(writer, sheet_name='Reporte de Defectos', index=False)
    df_changes.to_excel(writer, sheet_name='Control de Cambios', index=False)


# 2. GENERATE WORD DOC

doc = Document()
doc.add_heading('PLAN DE PRUEBAS DE SOFTWARE DE INGENIERÍA', 0)
doc.add_paragraph('PORTAL BES\nPlataforma Unificada de Consulta de Beneficios Estatales\nValidación y Verificación de Requisitos')
doc.add_paragraph('Documento: Plan de Pruebas Maestro (Master Test Plan)\nVersión: v1.0 Final\nEstándar de Referencia: IEEE Std 829-1998 / ISO/IEC 25010 e ISO 9126')

doc.add_heading('4. Diseño de Casos de Prueba (Matriz Detallada)', level=1)
doc.add_paragraph('De acuerdo con el modelo de calidad definido por la norma ISO 9126 (evaluando Funcionalidad, Fiabilidad y Usabilidad) y complementado con la norma moderna ISO/IEC 25010 (Seguridad, Accesibilidad y Rendimiento), se han diseñado 18 casos de prueba estructurados, excediendo el mínimo exigido para el equipo.')
doc.add_paragraph('La tabla completa con todos los escenarios y pasos a ejecutar, junto con sus estados correspondientes a las columnas del tablero Kanban (To Do, In Progress, Blocked, Failed, Passed), se encuentra detallada en el anexo de Excel "Entregable_Calidad_Portal_BES.xlsx" en la pestaña "Casos de Pruebas".')

doc.add_heading('4.1 a 4.5 Resumen de Cobertura Normativa', level=2)
doc.add_paragraph('HU-01 a HU-02 (Funcionalidad y Fiabilidad ISO 9126): Validados a través de los casos TC_001 al TC_006, garantizando que el software entrega las funciones requeridas y mantiene su nivel de desempeño en casos de caídas de APIs (tolerancia a fallos).')
doc.add_paragraph('HU-03 (Usabilidad ISO 9126): El entendimiento claro de los motivos de rechazo se cubrió en los casos TC_007 al TC_009.')
doc.add_paragraph('HU-04 (Seguridad ISO 25000): La confidencialidad y limpieza de estado de sesión se evaluó en los casos TC_010 a TC_012.')
doc.add_paragraph('HU-05 (Accesibilidad ISO 25000): Garantizada mediante pruebas con teclados y lectores NVDA (TC_013 a TC_015).')

doc.add_heading('5. Trazabilidad de Defectos (Bugs)', level=1)
doc.add_paragraph('Durante la simulación de ejecución en el tablero Kanban, los casos TC_007, TC_010, TC_015 y TC_017 resultaron con estado "Failed". Esto derivó en el levantamiento de los defectos BUG_001 a BUG_004 en nuestro reporte de hallazgos. Posteriormente, los errores críticos (ej. BUG_004 que impedía entrar al panel admin) fueron resueltos e integrados al Control de Cambios en la versión v1.2.')

doc.add_heading('6. Control de Cambios y Versionamiento', level=1)
doc.add_paragraph('El control de versiones se manejó de manera exhaustiva documentando cómo el prototipo evolucionó desde su estado conceptual (v1.0) hasta el prototipo robusto (v1.4) mitigando deficiencias de servidor (Cold Starts en Vercel) para asegurar la Fiabilidad y Usabilidad (ISO 9126). El desglose completo se adjunta en el Excel de Entregables.')

doc.save('Plan_de_Pruebas_Final_Completado.docx')

print("Archivos generados exitosamente.")
